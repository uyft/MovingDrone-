"""
DOCX 报告生成模块
使用 python-docx + matplotlib 生成带排版的中文 Word 报告
"""
import os
import math
import tempfile
from datetime import datetime, timezone, timedelta

import numpy as np

CST = timezone(timedelta(hours=8))  # 北京时间

from app.config import RESULT_DIR

# ── 中文字体 ──────────────────────────────────────────────────
_FONT_PATH = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
_FONT_PATH_BOLD = "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"

# ── 品牌色（十六进制，用于 matplotlib 图） ──────────────────
C_PRIMARY_HEX = "#143C78"
C_ACCENT_HEX  = "#0096C8"
C_DARK_HEX    = "#1E1E1E"
C_MEDIUM_HEX  = "#646464"


# ================================================================
#  主入口
# ================================================================
def build_task_report(task_id: str) -> str | None:
    from app.db import db_get_task, db_get_result

    task = db_get_task(task_id)
    result = db_get_result(task_id)
    if not task or not result:
        return None

    frames = result.get("frames", [])
    counts = [f.get("count", 0) for f in frames]
    peak_counts = [len(f.get("peaks", [])) for f in frames]

    n_frames = len(counts) or 1
    total_people = sum(counts)
    avg_count = total_people / n_frames
    max_count = max(counts) if counts else 0
    min_count = min(counts) if counts else 0
    peak_frame_idx = counts.index(max_count) + 1 if counts else 0
    std_count = float(np.std(counts)) if counts else 0.0
    cv = (std_count / avg_count * 100) if avg_count > 0 else 0.0
    density_label = _density_level(avg_count)
    trend = _trend_analysis(counts)
    high_threshold = avg_count + std_count
    high_density_ratio = sum(1 for c in counts if c > high_threshold) / n_frames * 100
    q25 = float(np.percentile(counts, 25)) if counts else 0.0
    q75 = float(np.percentile(counts, 75)) if counts else 0.0

    tmpdir = tempfile.mkdtemp(prefix="docx_report_")
    try:
        chart_path = _make_count_chart(counts, task["filename"], tmpdir)
        hist_path  = _make_density_histogram(counts, tmpdir)
        frame_imgs = _extract_key_frames(
            result.get("video_path", ""), frames, counts, peak_frame_idx, tmpdir)

        docx_path = os.path.join(RESULT_DIR, f"{task_id}_report.docx")
        _build_docx(
            docx_path=docx_path,
            task=task, result=result, n_frames=n_frames,
            total_people=total_people, avg_count=avg_count,
            max_count=max_count, min_count=min_count,
            peak_frame_idx=peak_frame_idx, std_count=std_count,
            cv=cv, density_label=density_label, trend=trend,
            high_density_ratio=high_density_ratio,
            q25=q25, q75=q75,
            counts=counts, peak_counts=peak_counts,
            chart_path=chart_path, hist_path=hist_path,
            frame_imgs=frame_imgs,
        )
        return docx_path
    except Exception:
        import traceback
        traceback.print_exc()
        return None
    finally:
        import shutil
        shutil.rmtree(tmpdir, ignore_errors=True)


# ================================================================
#  辅助分析函数
# ================================================================
def _density_level(avg: float) -> str:
    if avg < 15:   return "低密度（稀疏）"
    elif avg < 50:  return "中低密度"
    elif avg < 100: return "中等密度"
    elif avg < 300: return "中高密度"
    else:           return "高密度（密集）"


def _trend_analysis(counts: list[int]) -> dict:
    n = len(counts)
    if n < 10:
        return {"direction": "平稳", "desc": "视频帧数较少，无明显变化趋势。", "change_pct": 0}
    half = n // 2
    first_avg  = float(np.mean(counts[:half]))
    second_avg = float(np.mean(counts[half:]))
    change_pct = (second_avg - first_avg) / max(first_avg, 1) * 100

    if change_pct > 15:       direction, desc_extra = "明显上升", f"后半段平均人数（{second_avg:.1f}）较前半段（{first_avg:.1f}）增长 {change_pct:.0f}%，场景内人群数量呈上升趋势。"
    elif change_pct > 5:      direction, desc_extra = "小幅上升", f"后半段平均人数（{second_avg:.1f}）较前半段（{first_avg:.1f}）增长 {change_pct:.0f}%，人群数量略有增加。"
    elif change_pct > -5:     direction, desc_extra = "基本平稳", "视频前后段人群数量变化不大，整体保持稳定。"
    elif change_pct > -15:    direction, desc_extra = "小幅下降", f"后半段平均人数（{second_avg:.1f}）较前半段（{first_avg:.1f}）减少 {abs(change_pct):.0f}%，人群数量略有减少。"
    else:                     direction, desc_extra = "明显下降", f"后半段平均人数（{second_avg:.1f}）较前半段（{first_avg:.1f}）减少 {abs(change_pct):.0f}%，场景内人群数量呈下降趋势。"

    return {"direction": direction, "desc": desc_extra, "change_pct": change_pct}


def _model_label(name: str) -> str:
    return {
        "STEERER": "STEERER 密度估计模型",
        "YOLO11": "YOLO11 目标检测模型",
        "GD3A_VGG16": "GD³A + VGG16 检测模型",
        "GD3A_ResNet50": "GD³A + ResNet50 检测模型",
        "STNNet": "STNNet 时空网络模型",
    }.get(name, name)


# ================================================================
#  Matplotlib 图表
# ================================================================
def _make_count_chart(counts, title, tmpdir):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import font_manager
    font_manager.fontManager.addfont(_FONT_PATH)
    prop = font_manager.FontProperties(fname=_FONT_PATH)
    matplotlib.rcParams["font.family"] = prop.get_name()
    matplotlib.rcParams["axes.unicode_minus"] = False

    fig, ax = plt.subplots(figsize=(9, 3.6), dpi=150)
    xs = list(range(1, len(counts) + 1))
    avg = np.mean(counts) if counts else 0

    ax.plot(xs, counts, linewidth=1.2, color="#1478C8", alpha=0.92)
    ax.fill_between(xs, counts, alpha=0.10, color="#1478C8")
    ax.axhline(y=avg, color="#E05050", linewidth=0.9, linestyle="--", alpha=0.7)
    ax.text(len(counts) * 0.98, avg + max(counts) * 0.025, f"均值 {avg:.1f}",
            fontsize=8, color="#E05050", ha="right", va="bottom")

    if counts:
        max_i = int(np.argmax(counts))
        ax.annotate(f"峰值 {counts[max_i]}",
                    xy=(max_i + 1, counts[max_i]),
                    xytext=(max_i + 1 + len(counts) * 0.06, counts[max_i] * 1.08),
                    fontsize=8, color="#C85020",
                    arrowprops=dict(arrowstyle="->", color="#C85020", lw=0.9))

    ax.set_xlabel("帧序号", fontsize=10)
    ax.set_ylabel("人数", fontsize=10)
    ax.set_title(f"人群计数曲线 — {title}", fontsize=12, fontweight="bold")
    ax.set_xlim(1, len(counts) or 1)
    ax.tick_params(labelsize=8)
    ax.grid(True, linestyle=":", alpha=0.35)
    fig.tight_layout(pad=1.5)

    path = os.path.join(tmpdir, "count_chart.png")
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path


def _make_density_histogram(counts, tmpdir):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import font_manager
    font_manager.fontManager.addfont(_FONT_PATH)
    prop = font_manager.FontProperties(fname=_FONT_PATH)
    matplotlib.rcParams["font.family"] = prop.get_name()
    matplotlib.rcParams["axes.unicode_minus"] = False

    fig, ax = plt.subplots(figsize=(9, 3.2), dpi=130)
    bins = max(8, min(30, int(np.sqrt(len(counts))) * 2))
    ax.hist(counts, bins=bins, color="#1478C8", alpha=0.80, edgecolor="white", linewidth=0.5)
    ax.axvline(np.mean(counts), color="#E05050", linestyle="--", linewidth=1.1,
               label=f"均值 {np.mean(counts):.1f}")
    ax.set_xlabel("每帧人数", fontsize=10)
    ax.set_ylabel("帧数", fontsize=10)
    ax.set_title("人群数量分布直方图", fontsize=12, fontweight="bold")
    ax.legend(fontsize=8)
    ax.tick_params(labelsize=8)
    ax.grid(True, linestyle=":", alpha=0.35)
    fig.tight_layout(pad=1.5)

    path = os.path.join(tmpdir, "density_hist.png")
    fig.savefig(path, dpi=130, bbox_inches="tight")
    plt.close(fig)
    return path


# ================================================================
#  关键帧截图
# ================================================================
def _extract_key_frames(video_path, frames, counts, peak_idx, tmpdir):
    import cv2
    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        cap.release()
        return []

    total = int(cap.get(cv2.CAP_PROP_FRAME_COUNT)) or len(frames)
    indices = {peak_idx}
    for frac in [1, total // 3, total * 2 // 3, total]:
        indices.add(max(1, min(total, frac)))

    results = []
    for fi in sorted(indices):
        cap.set(cv2.CAP_PROP_POS_FRAMES, fi - 1)
        ret, frame_bgr = cap.read()
        if not ret:
            continue
        cnt = counts[fi - 1] if fi - 1 < len(counts) else 0
        h, w = frame_bgr.shape[:2]
        label = f"Frame #{fi}  |  Count: {cnt}"
        font_scale = max(0.6, w / 900)
        thickness = max(1, int(font_scale * 1.5))
        cv2.putText(frame_bgr, label, (12, max(28, int(h * 0.04))),
                    cv2.FONT_HERSHEY_SIMPLEX, font_scale, (0, 220, 255), thickness, cv2.LINE_AA)
        img_path = os.path.join(tmpdir, f"frame_{fi:06d}.jpg")
        cv2.imwrite(img_path, frame_bgr, [cv2.IMWRITE_JPEG_QUALITY, 80])
        results.append({"label": f"第 {fi} 帧", "path": img_path,
                        "frame_idx": fi, "count": cnt})
    cap.release()
    return results


# ================================================================
#  DOCX 构建
# ================================================================
def _build_docx(docx_path, task, result, n_frames, total_people, avg_count,
                max_count, min_count, peak_frame_idx, std_count, cv,
                density_label, trend, high_density_ratio, q25, q75,
                counts, peak_counts, chart_path, hist_path, frame_imgs):
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── 页面设置 ────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── 预设样式函数 ──────────────────────────────────────────
    def set_run_font(run, name="CJK", size=Pt(10.5), bold=False, color=C_DARK_HEX):
        run.font.name = name
        run.font.size = size
        run.bold = bold
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
        # 强制中文字体（东亚字体回退）
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), name)

    def add_heading_styled(text, level=1):
        """统一标题样式"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
        p.paragraph_format.space_after  = Pt(8)
        if level == 1:
            # 主标题
            run = p.add_run(text)
            set_run_font(run, "CJK", Pt(16), bold=True, color=C_PRIMARY_HEX)
        elif level == 2:
            run = p.add_run(text)
            set_run_font(run, "CJK", Pt(13), bold=True, color=C_DARK_HEX)

    def add_body(text, first_line_indent=True):
        """正文段落"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.line_spacing = Pt(20)
        if first_line_indent:
            p.paragraph_format.first_line_indent = Pt(21)  # 两个中文字
        run = p.add_run(text)
        set_run_font(run, "CJK", Pt(10.5), bold=False, color=C_DARK_HEX)

    def add_bullet(text, bold_prefix=""):
        """带可选粗体前缀的列表项"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.line_spacing = Pt(18)
        p.paragraph_format.left_indent  = Cm(0.6)
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            set_run_font(run_b, "CJK", Pt(10.5), bold=True, color=C_PRIMARY_HEX)
            run_t = p.add_run(text)
            set_run_font(run_t, "CJK", Pt(10.5), bold=False, color=C_DARK_HEX)
        else:
            run = p.add_run("• " + text)
            set_run_font(run, "CJK", Pt(10.5), bold=False, color=C_DARK_HEX)

    def add_image_centered(img_path, width_inches=5.8):
        """居中插入图片"""
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(width_inches))

    def set_cell_text(cell, text, bold=False, size=Pt(9), color=C_DARK_HEX, align="center"):
        """设置表格单元格文本并格式化"""
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER,
                       "left": WD_ALIGN_PARAGRAPH.LEFT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(str(text))
        set_run_font(run, "CJK", size, bold=bold, color=color)

    def shade_cells(row, color="F0F5FA"):
        """给一整行的所有单元格加背景色"""
        for cell in row.cells:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), color)
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)

    def add_thin_border(table):
        """给整个表格加细边框"""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            elem = OxmlElement(f'w:{edge}')
            elem.set(qn('w:val'), 'single')
            elem.set(qn('w:sz'), '4')
            elem.set(qn('w:space'), '0')
            elem.set(qn('w:color'), 'CCCCCC')
            borders.append(elem)
        tblPr.append(borders)


    # ================================================================
    #  封面
    # ================================================================
    for _ in range(4):
        doc.add_paragraph()  # 空行推到页面中部

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("DroneCrowd")
    set_run_font(run, "CJK", Pt(28), bold=True, color=C_PRIMARY_HEX)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("无人机人群智能感知平台")
    set_run_font(run, "CJK", Pt(14), bold=False, color=C_ACCENT_HEX)

    doc.add_paragraph()

    report_p = doc.add_paragraph()
    report_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = report_p.add_run("人群分析任务报告")
    set_run_font(run, "CJK", Pt(22), bold=True, color=C_DARK_HEX)

    doc.add_paragraph()
    doc.add_paragraph()

    # 元信息
    meta_items = [
        ("任务 ID", task['task_id']),
        ("视频文件", task.get('filename', '—')),
        ("分析模型", _model_label(task.get('model', ''))),
        ("分析模式", task.get('mode', '—')),
        ("生成时间", datetime.now(CST).strftime('%Y-%m-%d %H:%M:%S')),
    ]
    for label, val in meta_items:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp.paragraph_format.space_before = Pt(1)
        mp.paragraph_format.space_after  = Pt(1)
        r1 = mp.add_run(f"{label}：")
        set_run_font(r1, "CJK", Pt(10), bold=True, color=C_MEDIUM_HEX)
        r2 = mp.add_run(val)
        set_run_font(r2, "CJK", Pt(10), bold=False, color=C_DARK_HEX)

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("— 本报告由 DroneCrowd 平台自动生成 —")
    set_run_font(run, "CJK", Pt(9), bold=False, color=C_MEDIUM_HEX)

    doc.add_page_break()

    # ================================================================
    #  一、任务概览
    # ================================================================
    add_heading_styled("一、任务概览")

    # 概览表格（2列 × 4行）
    overview_data = [
        ("视频文件", task.get("filename", "—"), "分析模型", _model_label(task.get("model", ""))),
        ("分析模式", task.get("mode", "—"), "视频分辨率", f"{result.get('width', 0)} × {result.get('height', 0)}"),
        ("帧率 (FPS)", f"{result.get('fps', 0):.1f}", "总帧数", str(n_frames)),
        ("文件大小", f"{task.get('size_mb', 0):.1f} MB", "处理耗时", f"{result.get('total_time', 0):.1f} 秒"),
    ]
    tbl = doc.add_table(rows=len(overview_data), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_thin_border(tbl)

    for i, (l1, v1, l2, v2) in enumerate(overview_data):
        row = tbl.rows[i]
        set_cell_text(row.cells[0], l1, bold=True, size=Pt(9.5), color=C_MEDIUM_HEX, align="left")
        set_cell_text(row.cells[1], v1, bold=False, size=Pt(9.5), align="left")
        set_cell_text(row.cells[2], l2, bold=True, size=Pt(9.5), color=C_MEDIUM_HEX, align="left")
        set_cell_text(row.cells[3], v2, bold=False, size=Pt(9.5), align="left")
        # 设置列宽比例: 标签列窄，值列宽
        row.cells[0].width = Cm(2.8)
        row.cells[1].width = Cm(5.0)
        row.cells[2].width = Cm(2.8)
        row.cells[3].width = Cm(5.0)

    doc.add_paragraph()
    add_body(f"本报告基于 {_model_label(task.get('model', ''))} 对视频「{task.get('filename', '')}」"
             f"进行逐帧人群计数分析。视频共 {n_frames} 帧，分辨率 "
             f"{result.get('width', 0)}×{result.get('height', 0)}，帧率 {result.get('fps', 0):.0f} FPS，"
             f"总处理耗时 {result.get('total_time', 0):.1f} 秒。"
             f"以下各章节将从统计指标、时间变化趋势和空间分布三个维度对该视频的人群特征进行详细解读。")

    # ================================================================
    #  二、统计摘要
    # ================================================================
    add_heading_styled("二、统计摘要")

    stat_cards = [
        ("累计总人数", f"{total_people:,}", "平均每帧人数", f"{avg_count:.1f}"),
        ("最大单帧人数", f"{max_count}", "最小单帧人数", f"{min_count}"),
        ("峰值帧号", f"#{peak_frame_idx}", "标准差", f"{std_count:.1f}"),
        ("下四分位数 (Q1)", f"{q25:.0f}", "上四分位数 (Q3)", f"{q75:.0f}"),
    ]
    tbl2 = doc.add_table(rows=len(stat_cards), cols=4)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_thin_border(tbl2)

    for i, (l1, v1, l2, v2) in enumerate(stat_cards):
        row = tbl2.rows[i]
        shade_cells(row, "F0F5FA" if i % 2 == 0 else "FFFFFF")
        set_cell_text(row.cells[0], l1, bold=False, size=Pt(9.5), color=C_MEDIUM_HEX, align="left")
        set_cell_text(row.cells[1], v1, bold=True, size=Pt(11), color=C_PRIMARY_HEX)
        set_cell_text(row.cells[2], l2, bold=False, size=Pt(9.5), color=C_MEDIUM_HEX, align="left")
        set_cell_text(row.cells[3], v2, bold=True, size=Pt(11), color=C_PRIMARY_HEX)
        row.cells[0].width = Cm(3.2)
        row.cells[1].width = Cm(4.6)
        row.cells[2].width = Cm(3.2)
        row.cells[3].width = Cm(4.6)

    doc.add_paragraph()
    add_body(f"该视频的人群密度等级为「{density_label}」，每帧平均约 {avg_count:.1f} 人，"
             f"变异系数（CV）为 {cv:.1f}%，说明人群数量波动"
             f"{'较大' if cv > 40 else '适中' if cv > 20 else '较小'}。"
             f"单帧最高人数出现在第 {peak_frame_idx} 帧，达到 {max_count} 人；"
             f"最低为 {min_count} 人。标准差 {std_count:.1f}，"
             f"约 {high_density_ratio:.0f}% 的帧人数超过「均值+1σ」阈值，"
             f"这些高密度时段是需要重点关注的人群聚集期。")

    add_body(f"从四分位分布来看，25% 的帧人数低于 {q25:.0f}，75% 的帧人数低于 {q75:.0f}，"
             f"中间 50% 帧的人数集中在 [{q25:.0f}, {q75:.0f}] 区间内，"
             f"说明该场景中人群规模的整体波动范围约为 {q75 - q25:.0f} 人。")

    # ================================================================
    #  三、人群计数曲线
    # ================================================================
    doc.add_page_break()
    add_heading_styled("三、人群计数曲线")

    add_body(f"下图展示了逐帧人数变化趋势。蓝色实线为每帧检测到的人数，红色虚线为整体均值 "
             f"（{avg_count:.1f} 人/帧），橙色箭头标记了峰值位置。")

    add_image_centered(chart_path, 6.0)

    add_body(f"从曲线形态可以看出：{trend['desc']}")

    if max_count > avg_count * 2:
        add_body(f"需特别关注的是，峰值帧人数（{max_count} 人）超过均值的 2 倍，"
                 f"表明该时间点出现了瞬时人群聚集现象。建议结合下方关键帧截图查看该时刻的场景状态，"
                 f"判断是否存在异常事件或安全隐患。")

    # 密度分布直方图
    if os.path.exists(hist_path):
        doc.add_paragraph()
        add_body("下方为人数分布直方图，横轴为人数区间，纵轴为该区间内的帧数。"
                 "它直观展示了人群规模在不同区间的出现频率。")
        add_image_centered(hist_path, 6.0)

    # ================================================================
    #  四、关键帧截图
    # ================================================================
    if frame_imgs:
        doc.add_page_break()
        add_heading_styled("四、关键帧截图")

        add_body(f"以下为从视频中抽取的 {len(frame_imgs)} 个关键帧，"
                 f"包含首帧、末帧、1/3 处、2/3 处以及峰值帧。"
                 f"通过这些关键帧，可以直观对比不同时间点的人群分布情况。")

        # 2列网格布局用表格实现
        n = len(frame_imgs)
        for i in range(0, n, 2):
            img_tbl = doc.add_table(rows=1, cols=2)
            img_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j in range(2):
                idx = i + j
                cell = img_tbl.rows[0].cells[j]
                if idx < n:
                    fi = frame_imgs[idx]
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if os.path.exists(fi["path"]):
                        run = p.add_run()
                        run.add_picture(fi["path"], width=Inches(3.2))
                    # 标注
                    cap_p = cell.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap_p.add_run(f"{fi['label']}  |  人数: {fi['count']}")
                    set_run_font(cap_run, "CJK", Pt(9), bold=False, color=C_DARK_HEX)
                else:
                    set_cell_text(cell, "", size=Pt(9))
            doc.add_paragraph()

    # ================================================================
    #  五、帧统计附表
    # ================================================================
    doc.add_page_break()
    add_heading_styled("五、帧统计附表")

    add_body("下表为采样后的帧级统计数据（为避免表格过长，按比例间隔采样）。"
             "「人数」列为该帧检测到的总人数，「检测点」列为密度图中提取的局部峰值点数量。"
             "完整逐帧数据可通过平台的 JSON 导出功能获取。")

    n = len(counts)
    step = max(1, math.ceil(n / 40))
    sampled_indices = list(range(0, n, step))
    # 限制表格宽度，最多 40 行
    tbl3 = doc.add_table(rows=min(len(sampled_indices) + 1, 42), cols=4)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_thin_border(tbl3)

    # 表头
    hdr = tbl3.rows[0]
    shade_cells(hdr, C_PRIMARY_HEX.strip("#"))
    for j, h in enumerate(["序号", "帧号", "人数", "检测点"]):
        set_cell_text(hdr.cells[j], h, bold=True, size=Pt(9), color="FFFFFF")
        hdr.cells[j].width = Cm([2, 3, 3, 3][j])

    for row_i, fi in enumerate(sampled_indices):
        row = tbl3.rows[row_i + 1]
        if row_i % 2 == 0:
            shade_cells(row, "F5F7FA")
        for j, v in enumerate([str(row_i + 1), str(fi + 1), str(counts[fi]), str(peak_counts[fi])]):
            set_cell_text(row.cells[j], v, bold=False, size=Pt(8.5))

    # ================================================================
    #  六、总结与分析
    # ================================================================
    doc.add_page_break()
    add_heading_styled("六、总结与分析")

    add_body(f"本次分析基于 {_model_label(task.get('model', ''))} 对视频「{task.get('filename', '')}」"
             f"进行了全面的人群计数与时空分布分析。视频共计 {n_frames} 帧，"
             f"累计检测到 {total_people:,} 人次，属于{density_label}场景。以下是关键发现：")

    doc.add_paragraph()
    add_heading_styled("核心发现", level=2)

    findings = [
        (f"人群密度等级：", f"{density_label}，平均每帧 {avg_count:.1f} 人"),
        (f"变化趋势：", f"{trend['direction']}（{trend['change_pct']:+.0f}%），{trend['desc']}"),
        (f"波动程度：", f"标准差 {std_count:.1f}，变异系数 {cv:.1f}%——人群数量"
                      f"{'波动剧烈，存在明显的人员聚集和消散过程' if cv > 40 else '有一定波动，呈现自然的流动状态' if cv > 20 else '相对稳定，人群规模变化不大'}"),
        (f"峰值特征：", f"第 {peak_frame_idx} 帧达到峰值 {max_count} 人，"
                      f"该时刻为全视频人群最密集的时间点。建议重点关注该帧前后的场景变化。"),
        (f"高密度时段：", f"约 {high_density_ratio:.0f}% 的帧人数超过均值+1σ（{avg_count + std_count:.0f} 人），"
                        f"这些时段的人群密度需要特别关注。"),
        (f"人群分布：", f"中间 50% 帧的人数在 [{q25:.0f}, {q75:.0f}] 之间，"
                      f"最低 {min_count} 人，最高 {max_count} 人。"),
    ]
    for prefix, detail in findings:
        add_bullet(detail, bold_prefix=prefix)

    doc.add_paragraph()
    add_heading_styled("总体评价", level=2)
    add_body(_build_overall_assessment(density_label, trend, cv, high_density_ratio, max_count, avg_count))

    doc.add_paragraph()
    add_heading_styled("建议与展望", level=2)
    recommendations = _build_recommendations(density_label, cv, trend, max_count, avg_count)
    for i, rec in enumerate(recommendations, 1):
        add_body(f"{i}. {rec}", first_line_indent=True)

    # 结尾
    doc.add_paragraph()
    doc.add_paragraph()
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_p.add_run("— 报告结束 —")
    set_run_font(run, "CJK", Pt(10), bold=False, color=C_MEDIUM_HEX)

    end_p2 = doc.add_paragraph()
    end_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = end_p2.add_run(f"生成时间：{datetime.now(CST).strftime('%Y-%m-%d %H:%M:%S')}    "
                          "DroneCrowd 无人机人群智能感知平台")
    set_run_font(run2, "CJK", Pt(8), bold=False, color=C_MEDIUM_HEX)

    doc.save(docx_path)


# ================================================================
#  总结辅助
# ================================================================
def _build_overall_assessment(density_label, trend, cv, high_density_ratio, max_count, avg_count):
    parts = [f"综合来看，该视频呈现{density_label}的人群特征。"]

    if trend["direction"] in ("明显上升", "明显下降"):
        parts.append(f"人群数量在视频时间范围内{trend['direction']}，变化幅度较大，"
                    f"可能与场景中发生的事件（如集会、散场等）相关。")
    elif trend["direction"] in ("小幅上升", "小幅下降"):
        parts.append(f"人群数量{trend['direction']}，整体变化较为温和。")
    else:
        parts.append("人群数量整体保持平稳，未见明显的增减趋势。")

    if cv > 50:
        parts.append(f"人群波动剧烈（CV={cv:.0f}%），说明场景中存在明显的人员聚集与消散交替过程，"
                    f"可能对应特定事件周期。")
    elif cv > 25:
        parts.append(f"人群有一定波动（CV={cv:.0f}%），属于正常的人员流动范围。")
    else:
        parts.append(f"人群波动较小（CV={cv:.0f}%），密度分布相对均匀。")

    if max_count > avg_count * 3:
        parts.append(f"峰值人数远超均值，存在瞬时极端聚集情况，建议重点排查该时刻的安全风险。")
    elif max_count > avg_count * 2:
        parts.append(f"峰值人数明显高于均值，存在阶段性的人群聚集高峰。")

    return "".join(parts)


def _build_recommendations(density_label, cv, trend, max_count, avg_count):
    recs = []

    if "高密度" in density_label or "中高" in density_label:
        recs.append(
            "该场景人群密度较高，建议在峰值时段加强现场管控，设置分流通道和引导标识，"
            "避免因过度拥挤导致的安全隐患。")
    else:
        recs.append(
            "当前场景人群密度处于可控范围。建议保持常规监控，定期进行人数统计以建立"
            "正常基线，便于及时发现异常波动。")

    if cv > 40:
        recs.append(
            f"人群数量波动较大（CV={cv:.0f}%），建议关注波动背后的原因——是自然流动还是事件驱动。"
            f"可结合时间戳与现场日志，定位波动发生的具体时段，分析触发因素。")

    if trend["direction"] in ("明显上升", "明显下降"):
        recs.append(
            f"人群呈{trend['direction']}趋势，如该趋势持续，建议提前部署相应的人力与物资调配预案。")

    if max_count > avg_count * 2.5:
        recs.append(
            f"峰值人数（{max_count} 人）远超均值（{avg_count:.0f} 人），建议针对峰值时段设置预警阈值，"
            f"实现自动化告警，确保在人群快速聚集时能及时响应。")

    recs.append(
        "建议定期（如每周/每月）对同一场景进行重复分析，积累历史数据后可通过纵向对比"
        "发现人群模式的周期性变化，为资源调度提供数据支撑。")

    recs.append(
        "如需更详细的个体轨迹与运动模式分析，可使用平台的 Tracking 模式或 STNNet "
        "时空网络模型，获取每个人的运动轨迹，进一步分析人群流动方向与热点区域。")

    return recs
